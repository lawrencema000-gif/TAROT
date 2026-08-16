"""
Quick markdown -> .docx converter for MARKETING_STRATEGY.md.

Handles the subset of markdown actually used in the doc:
  - # / ## / ### / #### headings
  - Tables with | pipes
  - Bullet lists (-) and numbered lists (1.)
  - Bold (**), italic (*), inline code (`)
  - Blockquotes (>)
  - Horizontal rules (---)

Output is opinionated for Google Docs import: clean Title style,
no fancy theming, table borders set, font Calibri.

Run: python scripts/md-to-docx.py
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(__file__).parent.parent / "MARKETING_STRATEGY.md"
OUT = Path(__file__).parent.parent / "MARKETING_STRATEGY.docx"


def add_inline(paragraph, text):
    """Apply **bold**, *italic*, `code` runs to a paragraph."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("[") and "](" in token:
            link_text, _, rest = token[1:].partition("](")
            url = rest.rstrip(")")
            r = paragraph.add_run(link_text)
            r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            r.font.underline = True
            paragraph.add_run(f" ({url})").font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_cell_borders(cell):
    """Add solid borders to a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "C0C0C0")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def render_table(doc, lines):
    """Render a markdown table block to a docx table."""
    rows = [
        [c.strip() for c in line.strip().strip("|").split("|")]
        for line in lines
        if line.strip() and not re.match(r"^\|?[\s\-:|]+\|?\s*$", line)
    ]
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= n_cols:
                continue
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, cell_text)
            set_cell_borders(cell)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "F0F0F0")
                cell._tc.get_or_add_tcPr().append(shading_elm)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("# "):
            p = doc.add_heading(line.strip()[2:], level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            i += 1
            continue
        if line.strip().startswith("## "):
            doc.add_heading(line.strip()[3:], level=1)
            i += 1
            continue
        if line.strip().startswith("### "):
            doc.add_heading(line.strip()[4:], level=2)
            i += 1
            continue
        if line.strip().startswith("#### "):
            doc.add_heading(line.strip()[5:], level=3)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph().add_run("─" * 60).font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)
            i += 1
            continue

        if line.strip().startswith(">"):
            text_part = line.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(text_part)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # tables (start with `| `)
        if line.strip().startswith("|") and i + 1 < len(lines) and "|" in lines[i + 1]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            render_table(doc, table_lines)
            doc.add_paragraph()
            continue

        # numbered list
        m = re.match(r"^\s*(\d+)\.\s+(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            i += 1
            continue

        # bullet list
        m = re.match(r"^\s*-\s+(.+)", line)
        if m:
            indent = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5 + 0.3 * (indent // 2))
            add_inline(p, m.group(1))
            i += 1
            continue

        # checkbox-like list (- [ ])
        m = re.match(r"^\s*-\s*\[\s*\]\s*(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, "☐ " + m.group(1))
            i += 1
            continue

        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        # default: paragraph
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Size: {OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
